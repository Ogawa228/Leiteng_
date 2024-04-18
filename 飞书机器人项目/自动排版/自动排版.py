
import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QGridLayout, QWidget, QFileDialog, QMessageBox
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os

class App(QMainWindow):
    def __init__(self):
        super().__init__()
        self.title = '文档排版工具'
        self.left = 10
        self.top = 10
        self.width = 640
        self.height = 240
        self.input_file = ""
        self.output_folder = ""
        self.initUI()
    
    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)
        
        # 调整窗口大小以符合黄金比例
        width = 150
        golden_ratio = 1.618
        height = int(width / golden_ratio)
        self.resize(width, height)
        
        # Create central widget and layout
        widget = QWidget(self)
        self.setCentralWidget(widget)
        layout = QGridLayout(widget)

        # Buttons for input file, output folder, and triggering the modification
        self.inputButton = QPushButton('选择原文档', self)
        self.outputButton = QPushButton('选择输出文件夹', self)
        self.modifyButton = QPushButton('自动排版', self)
        
        self.inputButton.clicked.connect(self.openInputFileNameDialog)
        self.outputButton.clicked.connect(self.openOutputFolderDialog)
        self.modifyButton.clicked.connect(self.startModifyingDocument)

        # Place buttons in the layout
        layout.addWidget(self.inputButton, 1, 1)  # Second row, second column
        layout.addWidget(self.outputButton, 1, 3)  # Second row, third column
        layout.addWidget(self.modifyButton, 4, 2)  # Fourth row, fourth column

        self.show()
    
    def openInputFileNameDialog(self):
        options = QFileDialog.Options()
        self.input_file, _ = QFileDialog.getOpenFileName(self, "选择原文档", "", "Word Files (*.docx);;All Files (*)", options=options)
    
    def openOutputFolderDialog(self):
        options = QFileDialog.Options()
        self.output_folder = QFileDialog.getExistingDirectory(self, "选择输出文件夹", options=options)

    def startModifyingDocument(self):
            if self.input_file and self.output_folder:
                file_name = os.path.basename(self.input_file)
                output_path = os.path.join(self.output_folder, f"自动公文排版_{file_name}")
                self.modify_document(output_path)  # 注意这里只需要传入output_path
            else:
                QMessageBox.warning(self, "警告", "请先选择原文档和输出文件夹。")

    def startModifyingDocument(self):
        if self.input_file and self.output_folder:
            file_name = os.path.basename(self.input_file)
            output_path = os.path.join(self.output_folder, f"自动公文排版_{file_name}")
            self.modify_document(self.input_file, output_path)
        else:
            QMessageBox.warning(self, "警告", "请先选择原文档和输出文件夹。")

    def modify_document(self, input_path, output_path):
        doc = Document(input_path)
        
        # Page setup
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = Mm(37)
            section.bottom_margin = Mm(20)
            section.left_margin = Mm(28)
            section.right_margin = Mm(15)

        # Font settings
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith('题目部分'):
                font_name = '方正小标宋简体'
                font_size = Pt(18)  # Corresponds to font size "Small No. 2"
            elif paragraph.text.startswith('正文：'):
                font_name = '仿宋_GB2312'
                font_size = Pt(16)  # Corresponds to font size "No. 3"
                paragraph.paragraph_format.line_spacing = Pt(30)  # Line spacing of 30 pt
            elif any(paragraph.text.startswith(level) for level in ['一、', '二、', '三、']):
                font_name = '黑体'
                font_size = Pt(16)
            elif any(paragraph.text.startswith(level) for level in ['（一）', '（二）']):
                font_name = '楷体_GB2312'
                font_size = Pt(16)
            else:
                font_name = '仿宋_GB2312'
                font_size = Pt(16)

            for run in paragraph.runs:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = font_size
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # Save the document
        doc.save(output_path)
        QMessageBox.information(self, "完成", "文档排版完成！")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()
    sys.exit(app.exec_())