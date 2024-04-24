
#这是一个按照公文格式要求自动排版的工具。
#上传DOCX文件，工具将按照指定格式进行排版并输出一个新的DOCX文件。
import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QGridLayout, QWidget, QFileDialog, QMessageBox
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
from PyQt5.QtWidgets import QLabel
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt5.QtWidgets import QMessageBox


class App(QMainWindow):
    def __init__(self):
        super().__init__()
        self.title = '文档排版工具'
        self.left = 10
        self.top = 10
        self.width = 640
        self.height = 240
        self.input_file = ""
        self.output_folder = ""

        self.initUI()
    
    def initUI(self):
        self.setWindowTitle(self.title)
        # 窗口的初始位置，这里我们先随便设置一个，稍后会进行调整
        self.setGeometry(100, 100, self.width, self.height)

        # 添加应用说明的标签
        self.label_instructions = QLabel("这是一个按照公文格式要求自动排版的工具。上传DOCX文件，工具将按照指定格式进行排版并输出一个新的DOCX文件。", self)
        self.label_instructions.setWordWrap(True)  # 启用自动换行

        # 调整窗口大小以符合黄金比例或其他比例
        width = 400  # 可以根据实际情况调整这个宽度值
        golden_ratio = 3.618  # 黄金比例
        height = int(width / golden_ratio)
        self.resize(width, height)

        # Create central widget and layout
        widget = QWidget(self)
        self.setCentralWidget(widget)
        layout = QGridLayout(widget)

        # 应用说明的标签放置在布局中
        # 根据布局需要，可能需要调整行数，这里假设我们从第1行开始
        layout.addWidget(self.label_instructions, 0, 0, 1, 4)  # 第0行，占用4列

        # Buttons for input file, output folder, and triggering the modification
        self.inputButton = QPushButton('选择原文档', self)
        self.outputButton = QPushButton('选择输出文件夹', self)
        self.modifyButton = QPushButton('自动排版', self)

        self.inputButton.clicked.connect(self.openInputFileNameDialog)
        self.outputButton.clicked.connect(self.openOutputFolderDialog)
        self.modifyButton.clicked.connect(self.startModifyingDocument)

        # Place buttons in the layout
        # 根据布局需要，调整按钮的行位置
        layout.addWidget(self.inputButton, 2, 1)  # 第2行，第二列
        layout.addWidget(self.outputButton, 2, 2)  # 第2行，第三列
        layout.addWidget(self.modifyButton, 3, 1, 1, 2)  # 第3行，横跨两列

        # 将窗口移动到屏幕中心
        self.centerWindow()

        self.show()

    def centerWindow(self):
        # 获取屏幕的中心坐标
        screen = QApplication.desktop().screenGeometry()
        screenWidth = screen.width()
        screenHeight = screen.height()
        windowWidth = self.width
        windowHeight = self.height
        x = (screenWidth - windowWidth) // 2
        y = (screenHeight - windowHeight) // 2
        # 使用move方法将窗口移动到计算出的中心位置
        self.move(x, y)
    
    def openInputFileNameDialog(self):
        options = QFileDialog.Options()
        self.input_file, _ = QFileDialog.getOpenFileName(self, "选择原文档", "", "Word Files (*.docx);;All Files (*)", options=options)
    
    def openOutputFolderDialog(self):
        options = QFileDialog.Options()
        self.output_folder = QFileDialog.getExistingDirectory(self, "选择输出文件夹", options=options)

    def startModifyingDocument(self):
            if self.input_file and self.output_folder:
                file_name = os.path.basename(self.input_file)
                output_path = os.path.join(self.output_folder, f"自动公文排版_{file_name}")
                self.modify_document(output_path)  # 注意这里只需要传入output_path
            else:
                QMessageBox.warning(self, "警告", "请先选择原文档和输出文件夹。")

    def startModifyingDocument(self):
        if self.input_file and self.output_folder:
            file_name = os.path.basename(self.input_file)
            output_path = os.path.join(self.output_folder, f"自动公文排版_{file_name}")
            self.modify_document(self.input_file, output_path)
        else:
            QMessageBox.warning(self, "警告", "请先选择原文档和输出文件夹。")

    def modify_document(self, input_path, output_path):
        doc = Document(input_path)
        
        # 页面设置
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = Mm(37)
            section.bottom_margin = Mm(20)
            section.left_margin = Mm(28)
            section.right_margin = Mm(15)

        # 字体和段落格式设置
        for paragraph in doc.paragraphs:
            # 默认字体和大小设置
            font_name = '仿宋_GB2312'
            font_size = Pt(16)

            # 特殊格式处理
            if paragraph == doc.paragraphs[0]:  # 判断为文章大标题
                font_name = '方正小标宋简体'
                font_size = Pt(18)
            elif paragraph.text.startswith(('一、', '二、', '三、')):
                font_name = '黑体'
            elif paragraph.text.startswith(('（一）', '（二）')):
                font_name = '楷体_GB2312'

            # 段落首行缩进（排除标题）
            if not paragraph.text.startswith(('题目部分', '一、', '二、', '三、', '（一）', '（二）')):
                paragraph.paragraph_format.first_line_indent = Pt(16 * 2)

            for run in paragraph.runs:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = font_size
                # 设置西文字体以确保数字和特殊字符的字体
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

            paragraph.paragraph_format.line_spacing = Pt(30)  # 设置固定行间距为30磅
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        # 保存文档
        doc.save(output_path)
        QMessageBox.information(self, "完成", "文档排版完成！")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()
    sys.exit(app.exec_())